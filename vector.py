# vector.py
import os
import tempfile
import docx
from langchain.document_loaders import CSVLoader, PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain.vectorstores import FAISS
from langchain_core.documents import Document

def process_file(uploaded_file):
    """
    1) Save upload locally
    2) Load & chunk docs
    3) Build FAISS index (no SQLite dependency)
    4) Return (retriever, chunk_count)
    """
    # 1. Persist the upload
    path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # 2. Load raw documents
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext == ".csv":
        docs = CSVLoader(path).load()
    elif ext == ".pdf":
        docs = PyPDFLoader(path).load()
    else:  # .docx
        word = docx.Document(path)
        paras = [p.text for p in word.paragraphs if p.text.strip()]
        docs = [Document(page_content="\n\n".join(paras), metadata={})]

    # 3. Chunk text
    splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_documents(docs)
    n = len(chunks)

    # 4. Build FAISS vector store
    embeddings = OllamaEmbeddings(model="mxbai-embed-large")
    faiss_index = FAISS.from_documents(chunks, embeddings)

    retriever = faiss_index.as_retriever(search_kwargs={"k": 5})
    return retriever, n
